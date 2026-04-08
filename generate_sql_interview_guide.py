"""
Generate SQL Interview Questions & Answers for Data Analytics Roles
Outputs: .docx (MS Word) and .pdf
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from fpdf import FPDF
import textwrap

# ─────────────────────────────────────────────
# CONTENT: All questions organized by section
# ─────────────────────────────────────────────

sections = [
    # ============================================================
    # SECTION 1: SQL FUNDAMENTALS (Conceptual)
    # ============================================================
    {
        "title": "Section 1: SQL Fundamentals (Conceptual)",
        "questions": [
            {
                "q": "Q1. What is SQL and why is it important for Data Analytics?",
                "a": """SQL (Structured Query Language) is the standard language for managing and manipulating relational databases. For data analytics, SQL is critical because:

- It allows analysts to extract, filter, and aggregate large datasets directly from databases.
- It is the most widely used language for querying structured data in tools like MySQL, PostgreSQL, SQL Server, BigQuery, Snowflake, and Redshift.
- It enables repeatable, auditable analysis pipelines.
- Nearly every BI tool (Tableau, Power BI, Looker) uses SQL under the hood.

SQL is often the first and most important skill tested in data analytics interviews."""
            },
            {
                "q": "Q2. What is the difference between SQL and NoSQL databases?",
                "a": """SQL Databases (Relational):
- Store data in structured tables with rows and columns.
- Use a fixed schema (predefined structure).
- Support ACID transactions (Atomicity, Consistency, Isolation, Durability).
- Examples: MySQL, PostgreSQL, SQL Server, Oracle.

NoSQL Databases (Non-Relational):
- Store data in flexible formats: documents (MongoDB), key-value (Redis), wide-column (Cassandra), or graph (Neo4j).
- Schema-less or dynamic schema.
- Optimized for horizontal scaling and unstructured/semi-structured data.
- Examples: MongoDB, DynamoDB, Cassandra.

For analytics, SQL databases are preferred when data is structured and relationships matter. NoSQL is used for log data, real-time analytics, or unstructured data."""
            },
            {
                "q": "Q3. What are the different types of SQL commands? Explain with examples.",
                "a": """SQL commands are categorized into five types:

1. DDL (Data Definition Language) - Defines database structure
   - CREATE TABLE employees (id INT, name VARCHAR(100));
   - ALTER TABLE employees ADD COLUMN salary DECIMAL(10,2);
   - DROP TABLE employees;
   - TRUNCATE TABLE employees;

2. DML (Data Manipulation Language) - Manipulates data
   - INSERT INTO employees VALUES (1, 'Alice', 75000);
   - UPDATE employees SET salary = 80000 WHERE id = 1;
   - DELETE FROM employees WHERE id = 1;

3. DQL (Data Query Language) - Queries data
   - SELECT * FROM employees WHERE salary > 50000;

4. DCL (Data Control Language) - Controls access
   - GRANT SELECT ON employees TO analyst_role;
   - REVOKE INSERT ON employees FROM analyst_role;

5. TCL (Transaction Control Language) - Manages transactions
   - COMMIT;
   - ROLLBACK;
   - SAVEPOINT sp1;"""
            },
            {
                "q": "Q4. What is the difference between WHERE and HAVING?",
                "a": """WHERE filters rows BEFORE grouping (before GROUP BY is applied).
HAVING filters groups AFTER aggregation (after GROUP BY is applied).

Example:
-- WHERE: Filter individual rows
SELECT department, salary
FROM employees
WHERE salary > 50000;

-- HAVING: Filter aggregated groups
SELECT department, AVG(salary) AS avg_salary
FROM employees
GROUP BY department
HAVING AVG(salary) > 60000;

Key Rule: WHERE cannot use aggregate functions. HAVING can.

Common Interview Trap: "Can you use HAVING without GROUP BY?"
Answer: Technically yes (it treats the entire result as one group), but it is not standard practice."""
            },
            {
                "q": "Q5. What is the difference between DELETE, TRUNCATE, and DROP?",
                "a": """DELETE:
- DML command. Removes specific rows based on a condition.
- Can be rolled back (if inside a transaction).
- Fires triggers. Slower for large tables.
- Example: DELETE FROM orders WHERE order_date < '2023-01-01';

TRUNCATE:
- DDL command. Removes ALL rows from a table.
- Cannot be rolled back in most databases (except PostgreSQL).
- Does NOT fire triggers. Much faster than DELETE.
- Resets auto-increment counters.
- Example: TRUNCATE TABLE temp_logs;

DROP:
- DDL command. Removes the entire table (structure + data).
- Cannot be rolled back.
- Example: DROP TABLE temp_logs;

Analytics Context: Analysts rarely use DELETE/TRUNCATE/DROP directly. Understanding them matters for ETL pipelines and data warehouse management."""
            },
            {
                "q": "Q6. What is a Primary Key vs. a Foreign Key?",
                "a": """Primary Key:
- Uniquely identifies each row in a table.
- Cannot be NULL. Must be unique.
- A table can have only ONE primary key (can be composite).
- Example: employee_id in the employees table.

Foreign Key:
- A column that references the Primary Key of another table.
- Establishes a relationship between two tables.
- Can be NULL. Can have duplicates.
- Example: department_id in employees referencing id in departments.

CREATE TABLE departments (
    id INT PRIMARY KEY,
    name VARCHAR(100)
);

CREATE TABLE employees (
    id INT PRIMARY KEY,
    name VARCHAR(100),
    department_id INT,
    FOREIGN KEY (department_id) REFERENCES departments(id)
);

Analytics Relevance: Understanding keys is essential for writing correct JOINs and avoiding data duplication."""
            },
            {
                "q": "Q7. What is Normalization? Explain different Normal Forms.",
                "a": """Normalization is the process of organizing data to reduce redundancy and improve data integrity.

1NF (First Normal Form):
- Each column contains atomic (indivisible) values.
- Each row is unique.
- Bad: name = "Alice, Bob" -> Good: separate rows for Alice and Bob.

2NF (Second Normal Form):
- Must be in 1NF.
- No partial dependency (every non-key column depends on the ENTIRE primary key).
- Relevant when using composite primary keys.

3NF (Third Normal Form):
- Must be in 2NF.
- No transitive dependency (non-key columns should not depend on other non-key columns).
- Example: If employee table has department_name and department_location, remove department_location (it depends on department_name, not on employee_id).

BCNF (Boyce-Codd Normal Form):
- Stricter version of 3NF.
- Every determinant must be a candidate key.

Analytics Note: Data warehouses often use DENORMALIZED schemas (star/snowflake) for query performance. Analysts should understand normalization but will mostly work with denormalized data."""
            },
            {
                "q": "Q8. What is the difference between OLTP and OLAP?",
                "a": """OLTP (Online Transaction Processing):
- Designed for day-to-day transactions (inserts, updates, deletes).
- Highly normalized (3NF).
- Fast writes, short queries.
- Examples: Banking systems, e-commerce order processing.
- Databases: MySQL, PostgreSQL, SQL Server.

OLAP (Online Analytical Processing):
- Designed for complex analytical queries (aggregations, reporting).
- Denormalized (star/snowflake schema).
- Fast reads, complex queries over large datasets.
- Examples: Sales reporting, trend analysis, dashboards.
- Databases: BigQuery, Snowflake, Redshift, Clickhouse.

As a Data Analyst, you primarily work with OLAP systems, but you may query OLTP systems for real-time data."""
            },
        ]
    },

    # ============================================================
    # SECTION 2: JOINS & RELATIONSHIPS
    # ============================================================
    {
        "title": "Section 2: Joins & Relationships",
        "questions": [
            {
                "q": "Q9. Explain all types of SQL JOINs with examples.",
                "a": """Consider two tables:
employees: id | name | dept_id
departments: id | dept_name

1. INNER JOIN - Returns only matching rows from both tables.
SELECT e.name, d.dept_name
FROM employees e
INNER JOIN departments d ON e.dept_id = d.id;

2. LEFT JOIN (LEFT OUTER JOIN) - All rows from left table + matching rows from right. NULL if no match.
SELECT e.name, d.dept_name
FROM employees e
LEFT JOIN departments d ON e.dept_id = d.id;
-- Shows all employees, even those without a department.

3. RIGHT JOIN (RIGHT OUTER JOIN) - All rows from right table + matching rows from left.
SELECT e.name, d.dept_name
FROM employees e
RIGHT JOIN departments d ON e.dept_id = d.id;
-- Shows all departments, even those without employees.

4. FULL OUTER JOIN - All rows from both tables. NULL where no match.
SELECT e.name, d.dept_name
FROM employees e
FULL OUTER JOIN departments d ON e.dept_id = d.id;

5. CROSS JOIN - Cartesian product. Every row from table A paired with every row from table B.
SELECT e.name, d.dept_name
FROM employees e
CROSS JOIN departments d;

6. SELF JOIN - A table joined with itself.
SELECT e1.name AS employee, e2.name AS manager
FROM employees e1
JOIN employees e2 ON e1.manager_id = e2.id;

Interview Tip: Be ready to draw Venn diagrams for JOINs and explain when to use each."""
            },
            {
                "q": "Q10. What is the difference between JOIN and UNION?",
                "a": """JOIN combines columns from two or more tables horizontally (side by side) based on a related column.
UNION combines rows from two queries vertically (stacked on top) with the same column structure.

JOIN Example:
SELECT e.name, d.dept_name
FROM employees e
JOIN departments d ON e.dept_id = d.id;
-- Result has columns from BOTH tables.

UNION Example:
SELECT name, city FROM customers_us
UNION
SELECT name, city FROM customers_uk;
-- Result stacks rows from both queries.

UNION removes duplicates. UNION ALL keeps duplicates (faster).

Key Rules for UNION:
- Both queries must have the same number of columns.
- Corresponding columns must have compatible data types.
- Column names come from the first query."""
            },
            {
                "q": "Q11. Write a query to find employees who do NOT belong to any department.",
                "a": """Method 1: LEFT JOIN with NULL check (most common in interviews)
SELECT e.name
FROM employees e
LEFT JOIN departments d ON e.dept_id = d.id
WHERE d.id IS NULL;

Method 2: NOT IN
SELECT name
FROM employees
WHERE dept_id NOT IN (SELECT id FROM departments);
-- Warning: NOT IN fails silently if subquery returns NULL values.

Method 3: NOT EXISTS (safest and often most performant)
SELECT e.name
FROM employees e
WHERE NOT EXISTS (
    SELECT 1 FROM departments d WHERE d.id = e.dept_id
);

Analytics Tip: LEFT JOIN + IS NULL is the most readable and commonly expected answer. NOT EXISTS is the safest for production queries."""
            },
            {
                "q": "Q12. How do you find duplicate records in a table?",
                "a": """Method 1: GROUP BY + HAVING (most common)
SELECT email, COUNT(*) AS cnt
FROM customers
GROUP BY email
HAVING COUNT(*) > 1;

Method 2: Find full duplicate rows using ROW_NUMBER()
WITH ranked AS (
    SELECT *,
           ROW_NUMBER() OVER (PARTITION BY email ORDER BY id) AS rn
    FROM customers
)
SELECT * FROM ranked WHERE rn > 1;

Method 3: Self JOIN
SELECT a.*
FROM customers a
JOIN customers b
  ON a.email = b.email AND a.id > b.id;

To DELETE duplicates (keep one copy):
WITH ranked AS (
    SELECT id,
           ROW_NUMBER() OVER (PARTITION BY email ORDER BY id) AS rn
    FROM customers
)
DELETE FROM customers WHERE id IN (
    SELECT id FROM ranked WHERE rn > 1
);"""
            },
        ]
    },

    # ============================================================
    # SECTION 3: AGGREGATIONS & GROUP BY
    # ============================================================
    {
        "title": "Section 3: Aggregations & GROUP BY",
        "questions": [
            {
                "q": "Q13. Explain all aggregate functions in SQL.",
                "a": """Core Aggregate Functions:

COUNT(*) - Counts all rows (including NULLs)
COUNT(column) - Counts non-NULL values in a column
COUNT(DISTINCT column) - Counts unique non-NULL values

SUM(column) - Sum of all non-NULL values
AVG(column) - Average of all non-NULL values
MIN(column) - Minimum value
MAX(column) - Maximum value

Example:
SELECT
    department,
    COUNT(*) AS total_employees,
    COUNT(DISTINCT job_title) AS unique_titles,
    SUM(salary) AS total_payroll,
    AVG(salary) AS avg_salary,
    MIN(salary) AS min_salary,
    MAX(salary) AS max_salary
FROM employees
GROUP BY department;

Important NULL Behavior:
- COUNT(*) counts NULLs; COUNT(column) does NOT.
- SUM, AVG, MIN, MAX all ignore NULLs.
- AVG(salary) where values are 100, 200, NULL = 150 (not 100)."""
            },
            {
                "q": "Q14. What is the order of execution of a SQL query?",
                "a": """The logical order of execution (NOT the written order):

1. FROM / JOIN    - Tables are identified and joined
2. WHERE          - Rows are filtered (before grouping)
3. GROUP BY       - Rows are grouped
4. HAVING         - Groups are filtered (after aggregation)
5. SELECT         - Columns and expressions are computed
6. DISTINCT       - Duplicate rows are removed
7. ORDER BY       - Results are sorted
8. LIMIT / OFFSET - Results are paginated

Written Order:  SELECT ... FROM ... WHERE ... GROUP BY ... HAVING ... ORDER BY ... LIMIT ...
Execution Order: FROM -> WHERE -> GROUP BY -> HAVING -> SELECT -> ORDER BY -> LIMIT

Why This Matters:
- You cannot use a column alias defined in SELECT inside WHERE (it hasn't executed yet).
- You CAN use a column alias in ORDER BY (it executes after SELECT).
- HAVING can use aggregates because GROUP BY has already executed.

This is one of the most frequently asked conceptual questions in analytics interviews."""
            },
            {
                "q": "Q15. Write a query to find the second highest salary in each department.",
                "a": """Method 1: Using DENSE_RANK() (recommended)
WITH ranked AS (
    SELECT
        department,
        name,
        salary,
        DENSE_RANK() OVER (PARTITION BY department ORDER BY salary DESC) AS rnk
    FROM employees
)
SELECT department, name, salary
FROM ranked
WHERE rnk = 2;

Method 2: Using correlated subquery
SELECT e1.department, e1.name, e1.salary
FROM employees e1
WHERE 1 = (
    SELECT COUNT(DISTINCT e2.salary)
    FROM employees e2
    WHERE e2.department = e1.department
      AND e2.salary > e1.salary
);

Why DENSE_RANK() over ROW_NUMBER() or RANK():
- ROW_NUMBER(): If two people have the same highest salary, one gets rank 1 and the other rank 2. The "second highest" would actually be the same salary.
- RANK(): If two people tie for rank 1, the next rank is 3 (skips 2).
- DENSE_RANK(): If two people tie for rank 1, the next rank is 2. This correctly identifies the second distinct salary."""
            },
            {
                "q": "Q16. Write a query to calculate a running total of sales by date.",
                "a": """SELECT
    order_date,
    daily_sales,
    SUM(daily_sales) OVER (ORDER BY order_date) AS running_total
FROM (
    SELECT
        order_date,
        SUM(amount) AS daily_sales
    FROM orders
    GROUP BY order_date
) daily
ORDER BY order_date;

Output Example:
order_date  | daily_sales | running_total
2024-01-01  | 500         | 500
2024-01-02  | 300         | 800
2024-01-03  | 700         | 1500

For running total within each month:
SELECT
    order_date,
    SUM(amount) AS daily_sales,
    SUM(SUM(amount)) OVER (
        PARTITION BY DATE_TRUNC('month', order_date)
        ORDER BY order_date
    ) AS monthly_running_total
FROM orders
GROUP BY order_date
ORDER BY order_date;

This is a very common analytics interview question as running totals are fundamental to business reporting."""
            },
        ]
    },

    # ============================================================
    # SECTION 4: WINDOW FUNCTIONS (Advanced)
    # ============================================================
    {
        "title": "Section 4: Window Functions (Advanced)",
        "questions": [
            {
                "q": "Q17. What are Window Functions? How are they different from GROUP BY?",
                "a": """Window Functions perform calculations across a set of rows related to the current row WITHOUT collapsing rows into groups.

GROUP BY: Reduces multiple rows into one row per group.
Window Functions: Keep all individual rows while adding computed columns.

Example - GROUP BY:
SELECT department, AVG(salary) FROM employees GROUP BY department;
-- Returns one row per department.

Example - Window Function:
SELECT name, department, salary,
       AVG(salary) OVER (PARTITION BY department) AS dept_avg
FROM employees;
-- Returns ALL rows, each with its department average.

Types of Window Functions:

1. Ranking: ROW_NUMBER(), RANK(), DENSE_RANK(), NTILE()
2. Aggregate: SUM(), AVG(), COUNT(), MIN(), MAX() with OVER()
3. Value: LAG(), LEAD(), FIRST_VALUE(), LAST_VALUE(), NTH_VALUE()

Syntax:
function_name() OVER (
    PARTITION BY column    -- divides rows into groups
    ORDER BY column        -- defines row ordering within partition
    ROWS/RANGE BETWEEN ... -- defines the window frame
)"""
            },
            {
                "q": "Q18. Explain ROW_NUMBER(), RANK(), and DENSE_RANK() with an example.",
                "a": """Given data: Sales table with salesperson and revenue
| name  | revenue |
|-------|---------|
| Alice | 500     |
| Bob   | 500     |
| Carol | 400     |
| Dave  | 300     |

SELECT name, revenue,
    ROW_NUMBER() OVER (ORDER BY revenue DESC) AS row_num,
    RANK()       OVER (ORDER BY revenue DESC) AS rank,
    DENSE_RANK() OVER (ORDER BY revenue DESC) AS dense_rank
FROM sales;

Result:
| name  | revenue | row_num | rank | dense_rank |
|-------|---------|---------|------|------------|
| Alice | 500     | 1       | 1    | 1          |
| Bob   | 500     | 2       | 1    | 1          |
| Carol | 400     | 3       | 3    | 2          |
| Dave  | 300     | 4       | 4    | 3          |

- ROW_NUMBER(): Always unique. Ties get arbitrary ordering (1, 2).
- RANK(): Ties get the same rank. Next rank skips (1, 1, 3).
- DENSE_RANK(): Ties get the same rank. Next rank is consecutive (1, 1, 2).

When to use each:
- ROW_NUMBER(): Pagination, deduplication (keep one row per group).
- RANK(): Competitive ranking where gaps matter (sports standings).
- DENSE_RANK(): Finding Nth highest/lowest value."""
            },
            {
                "q": "Q19. Explain LAG() and LEAD() with a practical analytics example.",
                "a": """LAG() accesses a value from a PREVIOUS row.
LEAD() accesses a value from a NEXT row.

Syntax: LAG(column, offset, default) OVER (ORDER BY ...)

Example: Month-over-Month Revenue Growth
SELECT
    month,
    revenue,
    LAG(revenue, 1) OVER (ORDER BY month) AS prev_month_revenue,
    revenue - LAG(revenue, 1) OVER (ORDER BY month) AS mom_change,
    ROUND(
        (revenue - LAG(revenue, 1) OVER (ORDER BY month)) * 100.0
        / LAG(revenue, 1) OVER (ORDER BY month), 2
    ) AS mom_growth_pct
FROM monthly_revenue;

Result:
| month   | revenue | prev_month | mom_change | mom_growth_pct |
|---------|---------|------------|------------|----------------|
| 2024-01 | 10000   | NULL       | NULL       | NULL           |
| 2024-02 | 12000   | 10000      | 2000       | 20.00          |
| 2024-03 | 11500   | 12000      | -500       | -4.17          |

Example: Days Between User Sessions
SELECT
    user_id,
    session_date,
    LAG(session_date) OVER (PARTITION BY user_id ORDER BY session_date) AS prev_session,
    session_date - LAG(session_date) OVER (
        PARTITION BY user_id ORDER BY session_date
    ) AS days_between_sessions
FROM user_sessions;

LAG/LEAD are among the most commonly tested window functions in analytics interviews."""
            },
            {
                "q": "Q20. What are Window Frames (ROWS BETWEEN)? Explain with examples.",
                "a": """Window Frames define which rows within a partition are included in the calculation relative to the current row.

Syntax:
ROWS BETWEEN <start> AND <end>

Frame boundaries:
- UNBOUNDED PRECEDING: From the first row of the partition
- N PRECEDING: N rows before current row
- CURRENT ROW: The current row
- N FOLLOWING: N rows after current row
- UNBOUNDED FOLLOWING: To the last row of the partition

Example 1: 7-Day Moving Average
SELECT
    date,
    daily_revenue,
    AVG(daily_revenue) OVER (
        ORDER BY date
        ROWS BETWEEN 6 PRECEDING AND CURRENT ROW
    ) AS moving_avg_7d
FROM daily_sales;

Example 2: Running Total (default behavior)
SUM(amount) OVER (ORDER BY date ROWS BETWEEN UNBOUNDED PRECEDING AND CURRENT ROW)

Example 3: Centered Moving Average
AVG(value) OVER (ORDER BY date ROWS BETWEEN 3 PRECEDING AND 3 FOLLOWING)

Default Frame (when ORDER BY is specified):
RANGE BETWEEN UNBOUNDED PRECEDING AND CURRENT ROW
-- This means: from the start of the partition to the current row.

This is an advanced topic that separates strong candidates from average ones."""
            },
            {
                "q": "Q21. Write a query to calculate Year-over-Year growth for each product.",
                "a": """WITH yearly_sales AS (
    SELECT
        product_id,
        EXTRACT(YEAR FROM order_date) AS year,
        SUM(amount) AS annual_revenue
    FROM orders
    GROUP BY product_id, EXTRACT(YEAR FROM order_date)
)
SELECT
    product_id,
    year,
    annual_revenue,
    LAG(annual_revenue) OVER (
        PARTITION BY product_id ORDER BY year
    ) AS prev_year_revenue,
    ROUND(
        (annual_revenue - LAG(annual_revenue) OVER (
            PARTITION BY product_id ORDER BY year
        )) * 100.0
        / NULLIF(LAG(annual_revenue) OVER (
            PARTITION BY product_id ORDER BY year
        ), 0), 2
    ) AS yoy_growth_pct
FROM yearly_sales
ORDER BY product_id, year;

Key Points:
- NULLIF() prevents division by zero when previous year revenue is 0.
- PARTITION BY product_id ensures calculation is done per product.
- LAG() with PARTITION BY is the standard pattern for period-over-period analysis.

This pattern (with variations for QoQ, MoM, WoW) appears in almost every analytics interview."""
            },
        ]
    },

    # ============================================================
    # SECTION 5: SUBQUERIES & CTEs
    # ============================================================
    {
        "title": "Section 5: Subqueries & CTEs",
        "questions": [
            {
                "q": "Q22. What is a CTE? How is it different from a Subquery?",
                "a": """CTE (Common Table Expression) is a named temporary result set defined using the WITH clause. It exists only for the duration of the query.

CTE Example:
WITH high_earners AS (
    SELECT name, salary, department
    FROM employees
    WHERE salary > 100000
)
SELECT department, COUNT(*) AS count
FROM high_earners
GROUP BY department;

Equivalent Subquery:
SELECT department, COUNT(*) AS count
FROM (
    SELECT name, salary, department
    FROM employees
    WHERE salary > 100000
) AS high_earners
GROUP BY department;

Key Differences:
| Feature          | CTE                        | Subquery                   |
|------------------|----------------------------|----------------------------|
| Readability      | Much more readable         | Can get deeply nested      |
| Reusability      | Can reference multiple times| Must repeat the subquery   |
| Recursion        | Supports recursion         | No recursion               |
| Performance      | Usually same as subquery   | Usually same as CTE        |
| Scope            | Available for entire query | Available only where defined|

Best Practice: Use CTEs for readability. Use subqueries for simple, one-off filters."""
            },
            {
                "q": "Q23. What is a Correlated Subquery? Give an example.",
                "a": """A Correlated Subquery is a subquery that references columns from the outer query. It executes once for each row of the outer query.

Regular Subquery (executes once):
SELECT name FROM employees
WHERE salary > (SELECT AVG(salary) FROM employees);

Correlated Subquery (executes per row):
SELECT e1.name, e1.salary, e1.department
FROM employees e1
WHERE e1.salary > (
    SELECT AVG(e2.salary)
    FROM employees e2
    WHERE e2.department = e1.department  -- references outer query
);
-- Returns employees earning above their department average.

Another Example: Find latest order per customer
SELECT *
FROM orders o1
WHERE order_date = (
    SELECT MAX(o2.order_date)
    FROM orders o2
    WHERE o2.customer_id = o1.customer_id
);

Performance Note: Correlated subqueries can be slow on large datasets because they execute once per row. Window functions or JOINs are often more performant alternatives.

The window function equivalent:
WITH ranked AS (
    SELECT *, RANK() OVER (PARTITION BY customer_id ORDER BY order_date DESC) AS rn
    FROM orders
)
SELECT * FROM ranked WHERE rn = 1;"""
            },
            {
                "q": "Q24. Write a recursive CTE to generate a number sequence and explain how recursion works.",
                "a": """-- Generate numbers 1 to 10
WITH RECURSIVE numbers AS (
    -- Base case (anchor member)
    SELECT 1 AS n

    UNION ALL

    -- Recursive case
    SELECT n + 1
    FROM numbers
    WHERE n < 10
)
SELECT n FROM numbers;

How Recursive CTEs Work:
1. The anchor member executes first (SELECT 1 AS n) -> produces row: n=1.
2. The recursive member runs, referencing the CTE itself.
3. Each iteration uses the previous iteration's results.
4. Stops when the recursive member returns no new rows (WHERE n < 10 fails).

Practical Example: Organizational Hierarchy
WITH RECURSIVE org_tree AS (
    -- Base: CEO (no manager)
    SELECT id, name, manager_id, 1 AS level
    FROM employees
    WHERE manager_id IS NULL

    UNION ALL

    -- Recursive: Find direct reports
    SELECT e.id, e.name, e.manager_id, t.level + 1
    FROM employees e
    JOIN org_tree t ON e.manager_id = t.id
)
SELECT * FROM org_tree ORDER BY level, name;

Analytics Use Cases:
- Date series generation (fill gaps in time-series data).
- Hierarchical data traversal (org charts, category trees).
- Bill of materials explosion."""
            },
        ]
    },

    # ============================================================
    # SECTION 6: DATA TYPES, NULLS & TYPE CONVERSION
    # ============================================================
    {
        "title": "Section 6: Data Types, NULLs & Type Conversion",
        "questions": [
            {
                "q": "Q25. How does NULL work in SQL? What are common pitfalls?",
                "a": """NULL represents an unknown or missing value. It is NOT zero, empty string, or false.

Key Rules:
1. Any comparison with NULL returns NULL (not TRUE or FALSE).
   - NULL = NULL -> NULL (not TRUE!)
   - NULL != 1 -> NULL (not TRUE!)
   - Use IS NULL / IS NOT NULL for NULL checks.

2. NULL in arithmetic: Any operation with NULL returns NULL.
   - 5 + NULL = NULL
   - NULL * 0 = NULL

3. NULL in aggregates:
   - COUNT(*) counts NULLs; COUNT(column) ignores NULLs.
   - SUM, AVG, MIN, MAX all ignore NULLs.

4. NULL in WHERE: Rows where the condition evaluates to NULL are excluded.
   SELECT * FROM t WHERE col != 'A';  -- Excludes rows where col IS NULL!

Handling NULLs:
- COALESCE(col, 'default') - Returns first non-NULL value.
- NULLIF(a, b) - Returns NULL if a = b, else returns a.
- IFNULL(col, 0) (MySQL) or ISNULL(col, 0) (SQL Server).
- CASE WHEN col IS NULL THEN 'N/A' ELSE col END.

Interview Trap:
Q: What does "SELECT COUNT(*) FROM t WHERE col != 'A'" return if some rows have col = NULL?
A: It does NOT count NULL rows. To include them:
   WHERE col != 'A' OR col IS NULL"""
            },
            {
                "q": "Q26. What is the difference between CAST and CONVERT? When do you use COALESCE vs IFNULL?",
                "a": """CAST vs CONVERT:

CAST (ANSI Standard - works everywhere):
SELECT CAST('2024-01-15' AS DATE);
SELECT CAST(price AS DECIMAL(10,2));

CONVERT (SQL Server specific, has style parameter):
SELECT CONVERT(DATE, '2024-01-15');
SELECT CONVERT(VARCHAR, GETDATE(), 103);  -- dd/mm/yyyy format

Best Practice: Use CAST for portability. Use CONVERT only in SQL Server when you need specific formatting.

COALESCE vs IFNULL/ISNULL:

COALESCE (ANSI Standard - works everywhere):
- Accepts multiple arguments. Returns first non-NULL.
SELECT COALESCE(phone, mobile, email, 'No Contact') AS contact;

IFNULL (MySQL only):
- Accepts exactly two arguments.
SELECT IFNULL(phone, 'N/A');

ISNULL (SQL Server only):
- Accepts exactly two arguments.
- The return type is determined by the FIRST argument's type.
SELECT ISNULL(phone, 'N/A');

Best Practice: Use COALESCE for portability and flexibility."""
            },
        ]
    },

    # ============================================================
    # SECTION 7: STRING, DATE & CONDITIONAL FUNCTIONS
    # ============================================================
    {
        "title": "Section 7: String, Date & Conditional Functions",
        "questions": [
            {
                "q": "Q27. List commonly used string functions with examples.",
                "a": """-- Length
SELECT LENGTH('Hello');                    -- 5 (PostgreSQL/MySQL)
SELECT LEN('Hello');                       -- 5 (SQL Server)

-- Case conversion
SELECT UPPER('hello');                     -- HELLO
SELECT LOWER('HELLO');                     -- hello

-- Substring
SELECT SUBSTRING('Analytics', 1, 4);      -- Anal
SELECT LEFT('Analytics', 4);              -- Anal
SELECT RIGHT('Analytics', 4);             -- tics

-- Trim
SELECT TRIM('  hello  ');                 -- 'hello'
SELECT LTRIM('  hello');                  -- 'hello'
SELECT RTRIM('hello  ');                  -- 'hello'

-- Concatenation
SELECT CONCAT('Hello', ' ', 'World');     -- Hello World
SELECT 'Hello' || ' ' || 'World';        -- Hello World (PostgreSQL)

-- Replace
SELECT REPLACE('2024-01-15', '-', '/');   -- 2024/01/15

-- Position
SELECT POSITION('World' IN 'Hello World'); -- 7 (PostgreSQL)
SELECT CHARINDEX('World', 'Hello World');  -- 7 (SQL Server)

-- Pattern matching
SELECT * FROM users WHERE email LIKE '%@gmail.com';
SELECT * FROM users WHERE name LIKE 'A___';  -- A followed by 3 chars

-- Regex (PostgreSQL)
SELECT * FROM users WHERE email ~ '^[a-z]+@';"""
            },
            {
                "q": "Q28. Write common date function queries used in analytics.",
                "a": """-- Current date/time
SELECT CURRENT_DATE;                          -- 2024-01-15
SELECT CURRENT_TIMESTAMP;                     -- 2024-01-15 14:30:00
SELECT NOW();                                 -- Same as CURRENT_TIMESTAMP

-- Extract parts
SELECT EXTRACT(YEAR FROM order_date) AS yr;
SELECT EXTRACT(MONTH FROM order_date) AS mo;
SELECT EXTRACT(DOW FROM order_date) AS day_of_week;  -- 0=Sun (PostgreSQL)

-- Date truncation (essential for grouping)
SELECT DATE_TRUNC('month', order_date) AS month_start;
SELECT DATE_TRUNC('week', order_date) AS week_start;

-- Date arithmetic
SELECT order_date + INTERVAL '30 days';
SELECT order_date - INTERVAL '1 month';
SELECT AGE(NOW(), hire_date);                 -- PostgreSQL: interval

-- Date difference
SELECT DATEDIFF(day, start_date, end_date);   -- SQL Server
SELECT end_date - start_date;                 -- PostgreSQL (returns integer days)

-- Common Analytics Patterns:

-- Orders in last 30 days
SELECT * FROM orders WHERE order_date >= CURRENT_DATE - INTERVAL '30 days';

-- Group by month
SELECT DATE_TRUNC('month', order_date) AS month, SUM(amount)
FROM orders GROUP BY 1 ORDER BY 1;

-- Day-of-week analysis
SELECT
    TO_CHAR(order_date, 'Day') AS day_name,
    COUNT(*) AS order_count
FROM orders
GROUP BY EXTRACT(DOW FROM order_date), TO_CHAR(order_date, 'Day')
ORDER BY EXTRACT(DOW FROM order_date);"""
            },
            {
                "q": "Q29. Explain CASE WHEN with practical analytics examples.",
                "a": """CASE WHEN is SQL's IF-THEN-ELSE. Two forms:

Simple CASE:
SELECT name,
    CASE department
        WHEN 'Engineering' THEN 'Tech'
        WHEN 'Marketing' THEN 'Business'
        ELSE 'Other'
    END AS team
FROM employees;

Searched CASE (more flexible):
SELECT name, salary,
    CASE
        WHEN salary >= 150000 THEN 'Senior'
        WHEN salary >= 100000 THEN 'Mid'
        WHEN salary >= 50000 THEN 'Junior'
        ELSE 'Entry'
    END AS salary_band
FROM employees;

Analytics Use Cases:

1. Bucketing/Segmentation:
SELECT
    CASE
        WHEN age < 25 THEN '18-24'
        WHEN age < 35 THEN '25-34'
        WHEN age < 45 THEN '35-44'
        ELSE '45+'
    END AS age_group,
    COUNT(*) AS users
FROM users
GROUP BY 1;

2. Pivot-style aggregation:
SELECT
    product_id,
    SUM(CASE WHEN EXTRACT(QUARTER FROM order_date) = 1 THEN amount ELSE 0 END) AS q1_sales,
    SUM(CASE WHEN EXTRACT(QUARTER FROM order_date) = 2 THEN amount ELSE 0 END) AS q2_sales,
    SUM(CASE WHEN EXTRACT(QUARTER FROM order_date) = 3 THEN amount ELSE 0 END) AS q3_sales,
    SUM(CASE WHEN EXTRACT(QUARTER FROM order_date) = 4 THEN amount ELSE 0 END) AS q4_sales
FROM orders
GROUP BY product_id;

3. Conditional counting:
SELECT
    COUNT(*) AS total_orders,
    COUNT(CASE WHEN status = 'completed' THEN 1 END) AS completed,
    COUNT(CASE WHEN status = 'cancelled' THEN 1 END) AS cancelled
FROM orders;"""
            },
        ]
    },

    # ============================================================
    # SECTION 8: PERFORMANCE & OPTIMIZATION
    # ============================================================
    {
        "title": "Section 8: Performance & Optimization",
        "questions": [
            {
                "q": "Q30. What are Indexes? How do they improve query performance?",
                "a": """An Index is a data structure (usually B-tree) that speeds up data retrieval by creating a sorted reference to rows.

Without Index: Full table scan (reads every row) - O(n)
With Index: B-tree lookup - O(log n)

Creating Indexes:
-- Single column index
CREATE INDEX idx_email ON users(email);

-- Composite index (multi-column)
CREATE INDEX idx_dept_salary ON employees(department, salary);

-- Unique index
CREATE UNIQUE INDEX idx_unique_email ON users(email);

When Indexes Help:
- WHERE clause filters: WHERE email = 'alice@example.com'
- JOIN conditions: ON e.dept_id = d.id
- ORDER BY sorting
- GROUP BY grouping

When Indexes Hurt:
- Tables with very few rows (full scan is faster).
- Columns with low cardinality (e.g., gender with only M/F).
- Heavy INSERT/UPDATE/DELETE tables (index maintenance overhead).
- Wide indexes on rarely queried columns.

Composite Index Rule (Leftmost Prefix):
Index on (department, salary) helps:
  WHERE department = 'Sales'                     -- YES
  WHERE department = 'Sales' AND salary > 50000  -- YES
  WHERE salary > 50000                           -- NO (department is first)

Analytics Relevance: Analysts should understand indexes to write performant queries and communicate with database engineers about slow query optimization."""
            },
            {
                "q": "Q31. What is an Execution Plan? How do you use EXPLAIN?",
                "a": """An Execution Plan shows how the database engine will execute your query - which indexes it uses, join strategies, estimated costs, etc.

Usage:
-- PostgreSQL
EXPLAIN ANALYZE
SELECT e.name, d.dept_name
FROM employees e
JOIN departments d ON e.dept_id = d.id
WHERE e.salary > 50000;

-- MySQL
EXPLAIN
SELECT * FROM orders WHERE customer_id = 123;

Key Things to Look For:

1. Seq Scan (Sequential Scan): Full table scan. Usually bad for large tables.
   Fix: Add an index on the filtered column.

2. Index Scan / Index Only Scan: Using an index. Good.

3. Nested Loop vs Hash Join vs Merge Join:
   - Nested Loop: Good for small result sets.
   - Hash Join: Good for large, unsorted datasets.
   - Merge Join: Good for large, sorted datasets.

4. Sort: External sort can be expensive.
   Fix: Add index matching ORDER BY.

5. Estimated Rows vs Actual Rows: Large discrepancy means outdated statistics.
   Fix: ANALYZE table_name;

Common Optimization Tips:
- Avoid SELECT * (fetch only needed columns).
- Use EXISTS instead of IN for large subqueries.
- Avoid functions on indexed columns in WHERE (breaks index usage).
  Bad:  WHERE YEAR(order_date) = 2024
  Good: WHERE order_date >= '2024-01-01' AND order_date < '2025-01-01'"""
            },
            {
                "q": "Q32. What are Views, Materialized Views, and Temporary Tables?",
                "a": """VIEW: A saved query (virtual table). Does NOT store data. Re-executes the query each time.
CREATE VIEW active_customers AS
SELECT * FROM customers WHERE status = 'active';

-- Usage (just like a table)
SELECT * FROM active_customers WHERE city = 'NYC';

MATERIALIZED VIEW: A saved query that STORES the result physically. Must be refreshed manually.
CREATE MATERIALIZED VIEW monthly_sales AS
SELECT DATE_TRUNC('month', order_date) AS month, SUM(amount) AS total
FROM orders GROUP BY 1;

-- Refresh when underlying data changes
REFRESH MATERIALIZED VIEW monthly_sales;

TEMPORARY TABLE: A table that exists only for the current session.
CREATE TEMPORARY TABLE temp_results AS
SELECT * FROM orders WHERE amount > 1000;

Comparison:
| Feature           | View      | Materialized View | Temp Table   |
|-------------------|-----------|--------------------|--------------|
| Stores data       | No        | Yes                | Yes          |
| Auto-updates      | Yes       | No (manual refresh)| N/A          |
| Performance       | Same as query | Fast reads     | Fast reads   |
| Persistence       | Permanent | Permanent          | Session only |
| Use case          | Abstraction| Expensive queries  | Intermediate |

Analytics Best Practice:
- Views for reusable query logic and access control.
- Materialized Views for expensive aggregations in dashboards.
- Temp tables for multi-step transformations in ETL."""
            },
        ]
    },

    # ============================================================
    # SECTION 9: REAL-WORLD ANALYTICS SCENARIOS (Coding)
    # ============================================================
    {
        "title": "Section 9: Real-World Analytics Scenarios (Coding Questions)",
        "questions": [
            {
                "q": "Q33. Calculate user retention: What percentage of users who signed up in January were active in February?",
                "a": """WITH jan_signups AS (
    SELECT DISTINCT user_id
    FROM users
    WHERE DATE_TRUNC('month', signup_date) = '2024-01-01'
),
feb_active AS (
    SELECT DISTINCT user_id
    FROM user_activity
    WHERE activity_date >= '2024-02-01'
      AND activity_date < '2024-03-01'
)
SELECT
    COUNT(j.user_id) AS jan_signups,
    COUNT(f.user_id) AS retained_in_feb,
    ROUND(
        COUNT(f.user_id) * 100.0 / COUNT(j.user_id), 2
    ) AS retention_rate_pct
FROM jan_signups j
LEFT JOIN feb_active f ON j.user_id = f.user_id;

Generalized Cohort Retention (all months):
WITH cohorts AS (
    SELECT
        user_id,
        DATE_TRUNC('month', signup_date) AS cohort_month
    FROM users
),
activity AS (
    SELECT
        user_id,
        DATE_TRUNC('month', activity_date) AS activity_month
    FROM user_activity
)
SELECT
    c.cohort_month,
    EXTRACT(MONTH FROM AGE(a.activity_month, c.cohort_month)) AS months_since_signup,
    COUNT(DISTINCT a.user_id) AS active_users,
    COUNT(DISTINCT c.user_id) AS cohort_size,
    ROUND(COUNT(DISTINCT a.user_id) * 100.0 / COUNT(DISTINCT c.user_id), 2) AS retention_pct
FROM cohorts c
LEFT JOIN activity a ON c.user_id = a.user_id AND a.activity_month >= c.cohort_month
GROUP BY c.cohort_month, months_since_signup
ORDER BY c.cohort_month, months_since_signup;

This is one of the most important analytics interview questions. Cohort analysis is fundamental to product analytics."""
            },
            {
                "q": "Q34. Find the top 3 products by revenue in each category.",
                "a": """WITH product_revenue AS (
    SELECT
        p.category,
        p.product_name,
        SUM(o.quantity * o.unit_price) AS total_revenue
    FROM orders o
    JOIN products p ON o.product_id = p.id
    GROUP BY p.category, p.product_name
),
ranked AS (
    SELECT *,
        ROW_NUMBER() OVER (
            PARTITION BY category
            ORDER BY total_revenue DESC
        ) AS rn
    FROM product_revenue
)
SELECT category, product_name, total_revenue
FROM ranked
WHERE rn <= 3
ORDER BY category, rn;

Variation: If ties should be included (e.g., 3rd and 4th product have same revenue, show both):
-- Use DENSE_RANK() instead of ROW_NUMBER()

Variation: Show percentage of category revenue:
WITH product_revenue AS (
    SELECT
        p.category,
        p.product_name,
        SUM(o.quantity * o.unit_price) AS total_revenue
    FROM orders o
    JOIN products p ON o.product_id = p.id
    GROUP BY p.category, p.product_name
)
SELECT
    category,
    product_name,
    total_revenue,
    ROUND(total_revenue * 100.0 / SUM(total_revenue) OVER (PARTITION BY category), 2) AS pct_of_category
FROM product_revenue
ORDER BY category, total_revenue DESC;

This "Top N per group" pattern is extremely common in analytics interviews."""
            },
            {
                "q": "Q35. Find users who made purchases on 3 or more consecutive days.",
                "a": """WITH daily_purchases AS (
    SELECT DISTINCT user_id, DATE(order_date) AS purchase_date
    FROM orders
),
with_groups AS (
    SELECT
        user_id,
        purchase_date,
        purchase_date - ROW_NUMBER() OVER (
            PARTITION BY user_id ORDER BY purchase_date
        )::INT AS grp
    FROM daily_purchases
)
SELECT
    user_id,
    MIN(purchase_date) AS streak_start,
    MAX(purchase_date) AS streak_end,
    COUNT(*) AS consecutive_days
FROM with_groups
GROUP BY user_id, grp
HAVING COUNT(*) >= 3
ORDER BY consecutive_days DESC;

How This Works (Islands & Gaps technique):
Given dates: Jan 1, Jan 2, Jan 3, Jan 5
Row numbers:   1,     2,     3,     4
Date - RowNum: Dec 31, Dec 31, Dec 31, Jan 1

Consecutive dates produce the SAME (date - row_number) value, creating a "group."
Non-consecutive dates produce different values, creating a new group.

This "Islands & Gaps" pattern is a classic SQL interview problem that tests deep understanding of window functions."""
            },
            {
                "q": "Q36. Write a query to calculate a 7-day rolling average of daily active users (DAU).",
                "a": """WITH daily_active AS (
    SELECT
        DATE(activity_date) AS dt,
        COUNT(DISTINCT user_id) AS dau
    FROM user_activity
    GROUP BY DATE(activity_date)
)
SELECT
    dt,
    dau,
    ROUND(
        AVG(dau) OVER (
            ORDER BY dt
            ROWS BETWEEN 6 PRECEDING AND CURRENT ROW
        ), 2
    ) AS rolling_7d_avg
FROM daily_active
ORDER BY dt;

Important Details:
- ROWS BETWEEN 6 PRECEDING AND CURRENT ROW = 7 rows total (current + 6 prior).
- First 6 rows will have fewer than 7 days of data (partial window).
- To show NULL for incomplete windows:
  CASE
      WHEN ROW_NUMBER() OVER (ORDER BY dt) >= 7
      THEN AVG(dau) OVER (ORDER BY dt ROWS BETWEEN 6 PRECEDING AND CURRENT ROW)
  END AS rolling_7d_avg

Extending to 28-Day Rolling:
AVG(dau) OVER (ORDER BY dt ROWS BETWEEN 27 PRECEDING AND CURRENT ROW) AS rolling_28d_avg

This question tests:
1. Aggregation (COUNT DISTINCT for DAU)
2. Window functions with frame specification
3. Understanding of rolling calculations"""
            },
            {
                "q": "Q37. Identify churned users (users who were active last month but not this month).",
                "a": """WITH last_month_active AS (
    SELECT DISTINCT user_id
    FROM user_activity
    WHERE activity_date >= DATE_TRUNC('month', CURRENT_DATE) - INTERVAL '1 month'
      AND activity_date < DATE_TRUNC('month', CURRENT_DATE)
),
this_month_active AS (
    SELECT DISTINCT user_id
    FROM user_activity
    WHERE activity_date >= DATE_TRUNC('month', CURRENT_DATE)
      AND activity_date < DATE_TRUNC('month', CURRENT_DATE) + INTERVAL '1 month'
)
SELECT
    l.user_id AS churned_user
FROM last_month_active l
LEFT JOIN this_month_active t ON l.user_id = t.user_id
WHERE t.user_id IS NULL;

Full Status Classification:
WITH prev AS (
    SELECT DISTINCT user_id FROM user_activity
    WHERE activity_date >= '2024-02-01' AND activity_date < '2024-03-01'
),
curr AS (
    SELECT DISTINCT user_id FROM user_activity
    WHERE activity_date >= '2024-03-01' AND activity_date < '2024-04-01'
)
SELECT
    COALESCE(p.user_id, c.user_id) AS user_id,
    CASE
        WHEN p.user_id IS NOT NULL AND c.user_id IS NOT NULL THEN 'Retained'
        WHEN p.user_id IS NOT NULL AND c.user_id IS NULL THEN 'Churned'
        WHEN p.user_id IS NULL AND c.user_id IS NOT NULL THEN 'New/Reactivated'
    END AS status
FROM prev p
FULL OUTER JOIN curr c ON p.user_id = c.user_id;

This is critical for product analytics. Variants include:
- Churn rate calculation
- Churn prediction features
- Win-back campaign targeting"""
            },
            {
                "q": "Q38. Calculate the median salary per department (without using PERCENTILE_CONT).",
                "a": """Method 1: Using ROW_NUMBER (works in all databases)
WITH ranked AS (
    SELECT
        department,
        salary,
        ROW_NUMBER() OVER (PARTITION BY department ORDER BY salary) AS rn,
        COUNT(*) OVER (PARTITION BY department) AS cnt
    FROM employees
)
SELECT
    department,
    AVG(salary) AS median_salary
FROM ranked
WHERE rn IN (FLOOR((cnt + 1) / 2.0), CEIL((cnt + 1) / 2.0))
GROUP BY department;

How it works:
- For odd count (e.g., 5): FLOOR(3) = CEIL(3) = 3 -> middle value
- For even count (e.g., 6): FLOOR(3.5) = 3, CEIL(3.5) = 4 -> average of two middle values

Method 2: Using PERCENTILE_CONT (PostgreSQL, SQL Server)
SELECT
    department,
    PERCENTILE_CONT(0.5) WITHIN GROUP (ORDER BY salary) AS median_salary
FROM employees
GROUP BY department;

Method 3: Using NTILE
WITH quartiled AS (
    SELECT department, salary,
        NTILE(2) OVER (PARTITION BY department ORDER BY salary) AS half
    FROM employees
)
SELECT department,
    AVG(salary) AS approx_median
FROM quartiled
WHERE half = 1  -- approximation
GROUP BY department;

The manual method (Method 1) is the expected answer in interviews as it tests deeper understanding."""
            },
            {
                "q": "Q39. Write a query to create a funnel analysis: Signup -> Add to Cart -> Purchase.",
                "a": """WITH funnel AS (
    SELECT
        user_id,
        MAX(CASE WHEN event = 'signup' THEN 1 ELSE 0 END) AS signed_up,
        MAX(CASE WHEN event = 'add_to_cart' THEN 1 ELSE 0 END) AS added_to_cart,
        MAX(CASE WHEN event = 'purchase' THEN 1 ELSE 0 END) AS purchased
    FROM user_events
    WHERE event_date >= '2024-01-01' AND event_date < '2024-02-01'
    GROUP BY user_id
)
SELECT
    COUNT(*) AS total_users,
    SUM(signed_up) AS signups,
    SUM(added_to_cart) AS add_to_carts,
    SUM(purchased) AS purchases,
    ROUND(SUM(added_to_cart) * 100.0 / NULLIF(SUM(signed_up), 0), 2) AS signup_to_cart_pct,
    ROUND(SUM(purchased) * 100.0 / NULLIF(SUM(added_to_cart), 0), 2) AS cart_to_purchase_pct,
    ROUND(SUM(purchased) * 100.0 / NULLIF(SUM(signed_up), 0), 2) AS overall_conversion_pct
FROM funnel;

Ordered Funnel (events must happen in sequence):
WITH ordered_events AS (
    SELECT
        user_id,
        event,
        event_date,
        ROW_NUMBER() OVER (PARTITION BY user_id, event ORDER BY event_date) AS rn
    FROM user_events
),
funnel AS (
    SELECT
        s.user_id,
        s.event_date AS signup_date,
        c.event_date AS cart_date,
        p.event_date AS purchase_date
    FROM ordered_events s
    LEFT JOIN ordered_events c
        ON s.user_id = c.user_id AND c.event = 'add_to_cart'
        AND c.event_date > s.event_date AND c.rn = 1
    LEFT JOIN ordered_events p
        ON s.user_id = p.user_id AND p.event = 'purchase'
        AND p.event_date > c.event_date AND p.rn = 1
    WHERE s.event = 'signup' AND s.rn = 1
)
SELECT
    COUNT(*) AS signups,
    COUNT(cart_date) AS added_to_cart,
    COUNT(purchase_date) AS purchased,
    ROUND(COUNT(cart_date) * 100.0 / COUNT(*), 2) AS signup_to_cart,
    ROUND(COUNT(purchase_date) * 100.0 / NULLIF(COUNT(cart_date), 0), 2) AS cart_to_purchase
FROM funnel;

Funnel analysis is one of the most practical analytics interview topics."""
            },
            {
                "q": "Q40. Write a query to detect revenue anomalies (days where revenue deviated more than 2 standard deviations from the mean).",
                "a": """WITH daily_revenue AS (
    SELECT
        DATE(order_date) AS dt,
        SUM(amount) AS revenue
    FROM orders
    GROUP BY DATE(order_date)
),
stats AS (
    SELECT
        AVG(revenue) AS mean_revenue,
        STDDEV(revenue) AS stddev_revenue
    FROM daily_revenue
)
SELECT
    d.dt,
    d.revenue,
    s.mean_revenue,
    ROUND((d.revenue - s.mean_revenue) / s.stddev_revenue, 2) AS z_score,
    CASE
        WHEN d.revenue > s.mean_revenue + 2 * s.stddev_revenue THEN 'Spike'
        WHEN d.revenue < s.mean_revenue - 2 * s.stddev_revenue THEN 'Drop'
        ELSE 'Normal'
    END AS anomaly_status
FROM daily_revenue d
CROSS JOIN stats s
WHERE ABS(d.revenue - s.mean_revenue) > 2 * s.stddev_revenue
ORDER BY d.dt;

Rolling Anomaly Detection (30-day rolling stats):
WITH daily_revenue AS (
    SELECT DATE(order_date) AS dt, SUM(amount) AS revenue
    FROM orders GROUP BY 1
)
SELECT
    dt,
    revenue,
    AVG(revenue) OVER (ORDER BY dt ROWS BETWEEN 30 PRECEDING AND 1 PRECEDING) AS rolling_mean,
    STDDEV(revenue) OVER (ORDER BY dt ROWS BETWEEN 30 PRECEDING AND 1 PRECEDING) AS rolling_std,
    CASE
        WHEN revenue > AVG(revenue) OVER (ORDER BY dt ROWS BETWEEN 30 PRECEDING AND 1 PRECEDING)
            + 2 * STDDEV(revenue) OVER (ORDER BY dt ROWS BETWEEN 30 PRECEDING AND 1 PRECEDING)
        THEN 'Spike'
        WHEN revenue < AVG(revenue) OVER (ORDER BY dt ROWS BETWEEN 30 PRECEDING AND 1 PRECEDING)
            - 2 * STDDEV(revenue) OVER (ORDER BY dt ROWS BETWEEN 30 PRECEDING AND 1 PRECEDING)
        THEN 'Drop'
        ELSE 'Normal'
    END AS anomaly
FROM daily_revenue
ORDER BY dt;

This question tests statistical thinking combined with SQL - a key skill for analytics roles."""
            },
        ]
    },

    # ============================================================
    # SECTION 10: ADVANCED CONCEPTS & SCENARIO-BASED
    # ============================================================
    {
        "title": "Section 10: Advanced Concepts & Scenario-Based Questions",
        "questions": [
            {
                "q": "Q41. What is the difference between UNION, INTERSECT, and EXCEPT?",
                "a": """All three are set operations that combine results from two queries.

UNION: All rows from both queries (removes duplicates).
UNION ALL: All rows from both queries (keeps duplicates - faster).

SELECT city FROM customers_us
UNION
SELECT city FROM customers_uk;
-- Returns unique cities from both tables.

INTERSECT: Only rows that appear in BOTH queries.
SELECT user_id FROM jan_active
INTERSECT
SELECT user_id FROM feb_active;
-- Returns users active in BOTH January and February.

EXCEPT (MINUS in Oracle): Rows from first query that are NOT in second.
SELECT user_id FROM jan_active
EXCEPT
SELECT user_id FROM feb_active;
-- Returns users active in January but NOT in February (churned users).

Performance Tips:
- UNION ALL is faster than UNION (no deduplication step).
- INTERSECT and EXCEPT can often be rewritten as JOINs for better performance in some databases."""
            },
            {
                "q": "Q42. What are PIVOT and UNPIVOT? How do you pivot data in SQL?",
                "a": """PIVOT transforms rows into columns. UNPIVOT does the reverse.

Raw Data:
| product | quarter | revenue |
|---------|---------|---------|
| A       | Q1      | 100     |
| A       | Q2      | 150     |
| B       | Q1      | 200     |
| B       | Q2      | 250     |

Pivoted (using CASE WHEN - works in all databases):
SELECT
    product,
    SUM(CASE WHEN quarter = 'Q1' THEN revenue END) AS q1,
    SUM(CASE WHEN quarter = 'Q2' THEN revenue END) AS q2,
    SUM(CASE WHEN quarter = 'Q3' THEN revenue END) AS q3,
    SUM(CASE WHEN quarter = 'Q4' THEN revenue END) AS q4
FROM sales
GROUP BY product;

Result:
| product | q1  | q2  | q3   | q4   |
|---------|-----|-----|------|------|
| A       | 100 | 150 | NULL | NULL |
| B       | 200 | 250 | NULL | NULL |

SQL Server PIVOT syntax:
SELECT product, [Q1], [Q2], [Q3], [Q4]
FROM sales
PIVOT (SUM(revenue) FOR quarter IN ([Q1], [Q2], [Q3], [Q4])) AS pvt;

UNPIVOT (wide to long):
SELECT product, quarter, revenue
FROM pivoted_sales
UNPIVOT (revenue FOR quarter IN (q1, q2, q3, q4)) AS unpvt;

The CASE WHEN method is preferred in interviews as it works across all databases."""
            },
            {
                "q": "Q43. Explain the concept of a Star Schema vs. Snowflake Schema.",
                "a": """Both are data warehouse design patterns.

Star Schema:
- Central FACT table connected directly to DIMENSION tables.
- Dimensions are denormalized (flat).
- Simpler queries. Fewer JOINs. Faster reads.

         dim_product
              |
dim_date -- fact_sales -- dim_customer
              |
         dim_store

Snowflake Schema:
- Central FACT table connected to DIMENSION tables, which are further normalized into sub-dimensions.
- Dimensions are normalized.
- More JOINs. Less storage. More complex queries.

dim_date -- fact_sales -- dim_customer -- dim_city -- dim_country
              |
         dim_product -- dim_category -- dim_brand

Comparison:
| Feature       | Star Schema         | Snowflake Schema     |
|---------------|---------------------|----------------------|
| Complexity    | Simple              | Complex              |
| Joins         | Fewer               | More                 |
| Storage       | More (redundancy)   | Less (normalized)    |
| Query Speed   | Faster              | Slower               |
| Maintenance   | Easier              | Harder               |
| Best For      | BI/Reporting        | Strict storage needs |

Industry Standard: Star schema is overwhelmingly preferred for analytics and BI. Most modern data warehouses (Snowflake, BigQuery, Redshift) optimize for star schema patterns.

Key Vocabulary: Fact tables store measurable events (sales, clicks). Dimension tables store descriptive attributes (product name, customer segment, date)."""
            },
            {
                "q": "Q44. What is a Slowly Changing Dimension (SCD)? Explain Type 1, 2, and 3.",
                "a": """SCD handles dimensions that change over time (e.g., customer address, product price).

Type 1: Overwrite
- Simply update the old value with the new value.
- No history is preserved.
- Use when history doesn't matter.

Before: customer_id=1, city='NYC'
After:  customer_id=1, city='LA'

Type 2: Add New Row (most common in analytics)
- Add a new row with the updated value.
- Keep the old row with effective dates.
- Full history is preserved.

| customer_id | city | start_date | end_date   | is_current |
|-------------|------|------------|------------|------------|
| 1           | NYC  | 2020-01-01 | 2024-06-30 | FALSE      |
| 1           | LA   | 2024-07-01 | 9999-12-31 | TRUE       |

Type 3: Add New Column
- Add a column for the previous value.
- Only one level of history.

| customer_id | current_city | previous_city |
|-------------|-------------|---------------|
| 1           | LA          | NYC           |

Most Common in Practice:
- Type 1 for unimportant attributes.
- Type 2 for attributes that matter for historical analysis (most analytics use cases).
- Type 3 is rarely used.

Interview Tip: If asked about SCD, always mention Type 2 with effective dates and is_current flag. This is what analytics teams care about most."""
            },
            {
                "q": "Q45. Write a query to calculate Customer Lifetime Value (CLV) and segment customers.",
                "a": """WITH customer_metrics AS (
    SELECT
        c.customer_id,
        c.name,
        c.signup_date,
        COUNT(DISTINCT o.order_id) AS total_orders,
        SUM(o.amount) AS total_revenue,
        AVG(o.amount) AS avg_order_value,
        MIN(o.order_date) AS first_order,
        MAX(o.order_date) AS last_order,
        COUNT(DISTINCT DATE_TRUNC('month', o.order_date)) AS active_months,
        EXTRACT(EPOCH FROM MAX(o.order_date) - MIN(o.order_date)) / 86400 AS customer_lifespan_days
    FROM customers c
    LEFT JOIN orders o ON c.customer_id = o.customer_id
    GROUP BY c.customer_id, c.name, c.signup_date
),
clv AS (
    SELECT *,
        CASE
            WHEN customer_lifespan_days > 0
            THEN ROUND(
                (total_revenue / NULLIF(customer_lifespan_days, 0)) * 365, 2
            )
            ELSE total_revenue
        END AS estimated_annual_value,
        CASE
            WHEN active_months > 0
            THEN ROUND(total_revenue / active_months, 2)
            ELSE 0
        END AS avg_monthly_value
    FROM customer_metrics
)
SELECT
    customer_id,
    name,
    total_orders,
    total_revenue,
    avg_order_value,
    estimated_annual_value,
    CASE
        WHEN estimated_annual_value >= 10000 THEN 'Platinum'
        WHEN estimated_annual_value >= 5000 THEN 'Gold'
        WHEN estimated_annual_value >= 1000 THEN 'Silver'
        ELSE 'Bronze'
    END AS customer_segment,
    CASE
        WHEN last_order < CURRENT_DATE - INTERVAL '90 days' THEN 'At Risk'
        WHEN last_order < CURRENT_DATE - INTERVAL '180 days' THEN 'Churned'
        ELSE 'Active'
    END AS activity_status
FROM clv
ORDER BY estimated_annual_value DESC;

This combines multiple analytics concepts: aggregation, CASE WHEN, date functions, and business logic. It's a realistic analytics task that demonstrates practical SQL skills."""
            },
            {
                "q": "Q46. What are some common SQL query optimization techniques?",
                "a": """1. SELECT only needed columns (avoid SELECT *)
   Bad:  SELECT * FROM orders
   Good: SELECT order_id, amount, order_date FROM orders

2. Use WHERE to filter early
   Bad:  SELECT * FROM orders HAVING amount > 100
   Good: SELECT * FROM orders WHERE amount > 100

3. Avoid functions on indexed columns
   Bad:  WHERE YEAR(order_date) = 2024
   Good: WHERE order_date >= '2024-01-01' AND order_date < '2025-01-01'

4. Use EXISTS instead of IN for large subqueries
   Bad:  WHERE id IN (SELECT id FROM large_table)
   Good: WHERE EXISTS (SELECT 1 FROM large_table lt WHERE lt.id = t.id)

5. Use appropriate JOINs
   - INNER JOIN when you only need matching rows
   - Don't use LEFT JOIN if you then filter out NULLs in WHERE

6. Limit result sets
   - Use LIMIT during development
   - Use TOP/FETCH FIRST for production queries

7. Avoid DISTINCT when not needed (it requires sorting)

8. Use UNION ALL instead of UNION when duplicates are acceptable

9. Pre-aggregate in subqueries/CTEs before joining
   Bad:  JOIN on detail-level, then aggregate
   Good: Aggregate first, then JOIN summaries

10. Partition large tables by date for analytics queries

11. Use approximate functions for large datasets
    - APPROX_COUNT_DISTINCT instead of COUNT(DISTINCT) (BigQuery)
    - HyperLogLog for cardinality estimation

12. Avoid correlated subqueries on large tables (rewrite as JOINs or window functions)"""
            },
            {
                "q": "Q47. What is a Window Function PARTITION BY vs GROUP BY? Can you use both in the same query?",
                "a": """GROUP BY collapses rows: After GROUP BY, you get one row per group.
PARTITION BY keeps all rows: It defines the window for calculation but doesn't reduce rows.

Yes, you can use both in the same query!

Example: For each department, show the employee name, their salary, the department average (from GROUP BY), AND their rank within the department:

-- This WON'T work (can't show individual names with GROUP BY department):
SELECT name, department, AVG(salary) FROM employees GROUP BY department;

-- Correct: Use window function
SELECT
    name,
    department,
    salary,
    AVG(salary) OVER (PARTITION BY department) AS dept_avg,
    salary - AVG(salary) OVER (PARTITION BY department) AS diff_from_avg,
    RANK() OVER (PARTITION BY department ORDER BY salary DESC) AS salary_rank
FROM employees;

Using both together:
SELECT
    department,
    COUNT(*) AS emp_count,
    SUM(salary) AS total_salary,
    RANK() OVER (ORDER BY SUM(salary) DESC) AS dept_salary_rank
FROM employees
GROUP BY department;
-- GROUP BY aggregates per department.
-- RANK() ranks departments by their total salary.

Key Insight: Window functions execute AFTER GROUP BY in the query execution order, so they operate on the already-grouped result set."""
            },
            {
                "q": "Q48. Write a query to perform sessionization (group user events into sessions with a 30-minute inactivity timeout).",
                "a": """WITH events_with_gap AS (
    SELECT
        user_id,
        event_name,
        event_timestamp,
        LAG(event_timestamp) OVER (
            PARTITION BY user_id ORDER BY event_timestamp
        ) AS prev_event_time,
        EXTRACT(EPOCH FROM (
            event_timestamp - LAG(event_timestamp) OVER (
                PARTITION BY user_id ORDER BY event_timestamp
            )
        )) / 60 AS minutes_since_last_event
    FROM user_events
),
session_starts AS (
    SELECT *,
        CASE
            WHEN minutes_since_last_event IS NULL
              OR minutes_since_last_event > 30
            THEN 1
            ELSE 0
        END AS is_new_session
    FROM events_with_gap
),
sessions AS (
    SELECT *,
        SUM(is_new_session) OVER (
            PARTITION BY user_id ORDER BY event_timestamp
        ) AS session_id
    FROM session_starts
)
SELECT
    user_id,
    session_id,
    MIN(event_timestamp) AS session_start,
    MAX(event_timestamp) AS session_end,
    COUNT(*) AS events_in_session,
    EXTRACT(EPOCH FROM MAX(event_timestamp) - MIN(event_timestamp)) / 60 AS session_duration_minutes
FROM sessions
GROUP BY user_id, session_id
ORDER BY user_id, session_start;

How This Works:
1. LAG() finds time since previous event per user.
2. Flag rows where gap > 30 minutes as new session starts.
3. Running SUM of flags creates a session counter.
4. GROUP BY session_id to get session-level metrics.

This is an advanced question that tests:
- LAG() window function
- Conditional logic with CASE
- Cumulative SUM as a grouping technique
- Practical product analytics knowledge"""
            },
            {
                "q": "Q49. What are CTEs with INSERT/UPDATE/DELETE? (Writeable CTEs)",
                "a": """In PostgreSQL, CTEs can contain INSERT, UPDATE, or DELETE statements, with the results available to the main query.

Example 1: Delete duplicates and return what was deleted
WITH deleted AS (
    DELETE FROM customers
    WHERE id IN (
        SELECT id FROM (
            SELECT id,
                ROW_NUMBER() OVER (PARTITION BY email ORDER BY id) AS rn
            FROM customers
        ) ranked
        WHERE rn > 1
    )
    RETURNING *
)
SELECT COUNT(*) AS rows_deleted FROM deleted;

Example 2: Archive old orders
WITH archived AS (
    DELETE FROM orders
    WHERE order_date < '2020-01-01'
    RETURNING *
)
INSERT INTO orders_archive
SELECT * FROM archived;

Example 3: Upsert pattern
WITH new_data AS (
    SELECT 'alice@example.com' AS email, 'Alice' AS name
),
updated AS (
    UPDATE customers c
    SET name = n.name
    FROM new_data n
    WHERE c.email = n.email
    RETURNING c.email
)
INSERT INTO customers (email, name)
SELECT email, name FROM new_data
WHERE email NOT IN (SELECT email FROM updated);

Note: Writeable CTEs are PostgreSQL-specific. SQL Server and MySQL do not support them. This is an advanced topic that shows deep database knowledge."""
            },
            {
                "q": "Q50. Given a table of employee salaries, write a query to swap salaries between male and female employees without using a temporary table.",
                "a": """Table: employees (id, name, gender, salary)
Requirement: Swap so males get female salaries and vice versa (matched by some criteria, or simply swap the gender label).

Interpretation 1: Swap the gender label
UPDATE employees
SET gender = CASE
    WHEN gender = 'M' THEN 'F'
    WHEN gender = 'F' THEN 'M'
    ELSE gender
END;

Interpretation 2: Swap salaries between paired employees
-- Assuming employees are paired by rank within gender
WITH ranked AS (
    SELECT id, gender, salary,
        ROW_NUMBER() OVER (PARTITION BY gender ORDER BY id) AS rn
    FROM employees
),
pairs AS (
    SELECT
        m.id AS male_id, m.salary AS male_salary,
        f.id AS female_id, f.salary AS female_salary
    FROM ranked m
    JOIN ranked f ON m.rn = f.rn
    WHERE m.gender = 'M' AND f.gender = 'F'
)
UPDATE employees e
SET salary = CASE
    WHEN e.id IN (SELECT male_id FROM pairs)
        THEN (SELECT female_salary FROM pairs WHERE male_id = e.id)
    WHEN e.id IN (SELECT female_id FROM pairs)
        THEN (SELECT male_salary FROM pairs WHERE female_id = e.id)
    ELSE e.salary
END;

This is a popular tricky interview question. The key insight for Interpretation 1 is that CASE WHEN evaluates all conditions on the ORIGINAL values before applying any changes, so you don't need a temp variable.

Common Follow-up: "How does the database ensure both swaps happen atomically?"
Answer: SQL UPDATE statements are atomic - all SET expressions are evaluated using the original row values before any updates are applied."""
            },
        ]
    },
]

# Quick Reference Section
quick_reference = """
QUICK REFERENCE CHEAT SHEET

SQL Order of Execution:
FROM -> WHERE -> GROUP BY -> HAVING -> SELECT -> DISTINCT -> ORDER BY -> LIMIT

JOIN Types:
INNER JOIN  = Only matching rows
LEFT JOIN   = All left + matching right
RIGHT JOIN  = All right + matching left
FULL JOIN   = All from both
CROSS JOIN  = Cartesian product

Window Functions Syntax:
function() OVER (PARTITION BY col ORDER BY col ROWS BETWEEN ... AND ...)

Ranking Functions:
ROW_NUMBER() = Always unique (1,2,3,4)
RANK()       = Ties same, gaps (1,1,3,4)
DENSE_RANK() = Ties same, no gaps (1,1,2,3)

NULL Rules:
- NULL = NULL is NULL (not TRUE)
- COUNT(*) counts NULLs; COUNT(col) skips NULLs
- Use IS NULL / IS NOT NULL (not = NULL)
- COALESCE(a, b, c) returns first non-NULL

Common Analytics Patterns:
1. Top N per group     -> ROW_NUMBER() + PARTITION BY
2. Running total       -> SUM() OVER (ORDER BY date)
3. Period comparison   -> LAG() with PARTITION BY
4. Consecutive days    -> Islands & Gaps (date - ROW_NUMBER())
5. Funnel analysis     -> CASE WHEN pivoting
6. Cohort retention    -> Self-join on signup month vs activity month
7. Sessionization      -> LAG() + cumulative SUM of gap flags
8. Deduplication       -> ROW_NUMBER() PARTITION BY duplicate_cols
"""


# ─────────────────────────────────────────────
# GENERATE DOCX
# ─────────────────────────────────────────────

def generate_docx(filename):
    doc = Document()

    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    title = doc.add_heading('SQL Interview Questions & Answers', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Comprehensive Guide for Data Analytics Roles\nBasic to Advanced | Conceptual to Coding\n50 Questions with Detailed Answers')
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()  # spacer

    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    for sec in sections:
        toc_para = doc.add_paragraph(sec['title'], style='List Bullet')
        toc_para.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # Sections
    for sec in sections:
        doc.add_heading(sec['title'], level=1)

        for qa in sec['questions']:
            # Question
            q_para = doc.add_heading(qa['q'], level=2)

            # Answer
            doc.add_heading('Answer:', level=3)
            for line in qa['a'].strip().split('\n'):
                p = doc.add_paragraph(line)
                p.paragraph_format.space_after = Pt(2)
                # Style SQL code lines
                stripped = line.strip()
                if stripped.startswith(('SELECT', 'FROM', 'WHERE', 'JOIN', 'GROUP',
                    'ORDER', 'HAVING', 'INSERT', 'UPDATE', 'DELETE', 'CREATE',
                    'ALTER', 'DROP', 'WITH', 'GRANT', 'REVOKE', '--', 'UNION',
                    'INTERSECT', 'EXCEPT', 'TRUNCATE', 'EXPLAIN', 'REFRESH',
                    'RETURNING', 'SET ')):
                    for run in p.runs:
                        run.font.name = 'Consolas'
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(0, 100, 0)

            doc.add_paragraph()  # spacer

        doc.add_page_break()

    # Quick Reference
    doc.add_heading('Quick Reference Cheat Sheet', level=1)
    for line in quick_reference.strip().split('\n'):
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(1)

    doc.save(filename)
    print(f"DOCX saved: {filename}")


# ─────────────────────────────────────────────
# GENERATE PDF
# ─────────────────────────────────────────────

class PDFReport(FPDF):
    def header(self):
        self.set_font('Helvetica', 'I', 8)
        self.set_text_color(150, 150, 150)
        self.cell(0, 10, 'SQL Interview Guide for Data Analytics | Basic to Advanced', 0, 0, 'C')
        self.ln(10)

    def footer(self):
        self.set_y(-15)
        self.set_font('Helvetica', 'I', 8)
        self.set_text_color(150, 150, 150)
        self.cell(0, 10, f'Page {self.page_no()}/{{nb}}', 0, 0, 'C')

    def section_title(self, text):
        self.set_font('Helvetica', 'B', 16)
        self.set_text_color(25, 25, 112)  # dark blue
        self.multi_cell(0, 10, text)
        self.ln(4)

    def question_title(self, text):
        self.set_font('Helvetica', 'B', 12)
        self.set_text_color(0, 0, 0)
        self.multi_cell(0, 7, text)
        self.ln(2)

    def answer_label(self):
        self.set_font('Helvetica', 'B', 11)
        self.set_text_color(0, 100, 0)
        self.cell(0, 7, 'Answer:', 0, 1)
        self.ln(1)

    def answer_text(self, text):
        self.set_font('Helvetica', '', 10)
        self.set_text_color(30, 30, 30)
        # Handle encoding
        safe_text = text.encode('latin-1', 'replace').decode('latin-1')
        self.multi_cell(0, 5, safe_text)
        self.ln(3)


def generate_pdf(filename):
    pdf = PDFReport()
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=20)

    # Title Page
    pdf.add_page()
    pdf.ln(40)
    pdf.set_font('Helvetica', 'B', 28)
    pdf.set_text_color(25, 25, 112)
    pdf.cell(0, 15, 'SQL Interview Questions', 0, 1, 'C')
    pdf.cell(0, 15, '& Answers', 0, 1, 'C')
    pdf.ln(10)
    pdf.set_font('Helvetica', '', 16)
    pdf.set_text_color(80, 80, 80)
    pdf.cell(0, 10, 'Comprehensive Guide for Data Analytics Roles', 0, 1, 'C')
    pdf.cell(0, 10, 'Basic to Advanced | Conceptual to Coding', 0, 1, 'C')
    pdf.cell(0, 10, '50 Questions with Detailed Answers', 0, 1, 'C')
    pdf.ln(20)
    pdf.set_font('Helvetica', 'I', 12)
    pdf.set_text_color(120, 120, 120)
    pdf.cell(0, 10, 'Covers: Fundamentals, JOINs, Aggregations, Window Functions,', 0, 1, 'C')
    pdf.cell(0, 10, 'CTEs, Optimization, Real-World Scenarios & More', 0, 1, 'C')

    # Table of Contents
    pdf.add_page()
    pdf.set_font('Helvetica', 'B', 18)
    pdf.set_text_color(25, 25, 112)
    pdf.cell(0, 12, 'Table of Contents', 0, 1)
    pdf.ln(5)
    for i, sec in enumerate(sections, 1):
        pdf.set_font('Helvetica', '', 12)
        pdf.set_text_color(30, 30, 30)
        pdf.cell(0, 8, f"  {sec['title']}", 0, 1)

    # Content
    for sec in sections:
        pdf.add_page()
        pdf.section_title(sec['title'])

        for qa in sec['questions']:
            # Check if we need a new page (leave room for question + some answer)
            if pdf.get_y() > 240:
                pdf.add_page()

            pdf.question_title(qa['q'])
            pdf.answer_label()
            pdf.answer_text(qa['a'])

            # Separator line
            pdf.set_draw_color(200, 200, 200)
            pdf.line(10, pdf.get_y(), 200, pdf.get_y())
            pdf.ln(5)

    # Quick Reference
    pdf.add_page()
    pdf.section_title('Quick Reference Cheat Sheet')
    pdf.set_font('Courier', '', 9)
    pdf.set_text_color(30, 30, 30)
    safe_ref = quick_reference.encode('latin-1', 'replace').decode('latin-1')
    pdf.multi_cell(0, 4.5, safe_ref)

    pdf.output(filename)
    print(f"PDF saved: {filename}")


# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────

if __name__ == '__main__':
    output_dir = '/home/raethteam/upanshu/research/'

    docx_path = output_dir + 'SQL_Interview_Guide_Data_Analytics.docx'
    pdf_path = output_dir + 'SQL_Interview_Guide_Data_Analytics.pdf'

    generate_docx(docx_path)
    generate_pdf(pdf_path)

    # Summary
    total_q = sum(len(s['questions']) for s in sections)
    print(f"\nTotal Questions: {total_q}")
    print(f"Total Sections: {len(sections)}")
    print("Done!")
